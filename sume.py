import requests
import json
import os
import time
import re
from datetime import datetime, timedelta
from docx import Document
from ebooklib import epub
from openai import OpenAI

# ============================================================
# CẤU HÌNH
# ============================================================
api_key_path = r"D:\Code\API KEY\Chatgpt API Key.txt"
EMAIL_BASE_DIR = r"D:\gmail_emails"          # thư mục gốc chứa các folder ngày
OUTPUT_BASE_DIR = r"D:\Articles AI-Sum\Summaries"

try:
    with open(api_key_path, "r", encoding="utf-8") as key_file:
        api_key = key_file.read().strip()
except FileNotFoundError:
    print(f"Không tìm thấy file API key tại {api_key_path}")
    exit(1)

client = OpenAI(api_key=api_key)

# ============================================================
# HELPER FUNCTIONS
# ============================================================

def get_date_folders(n: int) -> list[str]:
    """Trả về danh sách đường dẫn folder của n ngày gần nhất (tồn tại trên disk)."""
    today = datetime.now()
    folders = []
    for i in range(n):
        day = today - timedelta(days=i)
        folder_name = day.strftime("%Y%m%d")
        folder_path = os.path.join(EMAIL_BASE_DIR, folder_name)
        if os.path.isdir(folder_path):
            folders.append(folder_path)
        else:
            print(f"[WARN] Không tìm thấy folder: {folder_path}")
    return folders


def get_email_files(folders: list[str]) -> list[tuple[str, str]]:
    """
    Quét tất cả file trong danh sách folder.
    Trả về list (full_path, filename_no_ext).
    Hỗ trợ mọi extension (.txt, .eml, .html, v.v.)
    """
    files = []
    for folder in folders:
        for fname in os.listdir(folder):
            fpath = os.path.join(folder, fname)
            if os.path.isfile(fpath):
                name_no_ext = os.path.splitext(fname)[0]
                files.append((fpath, name_no_ext))
    return files


def match_keywords(name_no_ext: str, keywords: list[str]) -> bool:
    """Kiểm tra xem tên file có chứa bất kỳ keyword nào không (case-insensitive)."""
    name_lower = name_no_ext.lower()
    return any(kw.lower() in name_lower for kw in keywords)


def filter_emails(
    all_files: list[tuple[str, str]],
    include_kws: list[str],
    exclude_kws: list[str]
) -> list[tuple[str, str]]:
    """
    Lọc email theo include/exclude keywords.
    - include_kws rỗng  → lấy tất cả
    - exclude_kws rỗng  → không loại bỏ gì
    """
    result = []
    for fpath, name in all_files:
        # Bước 1: include filter
        if include_kws:
            if not match_keywords(name, include_kws):
                continue
        # Bước 2: exclude filter
        if exclude_kws:
            if match_keywords(name, exclude_kws):
                continue
        result.append((fpath, name))
    return result


def read_email_content(fpath: str) -> str | None:
    """Đọc nội dung file email (txt/eml/html). Trả về None nếu lỗi."""
    try:
        with open(fpath, "r", encoding="utf-8", errors="replace") as f:
            return f.read()
    except Exception as e:
        print(f"[ERROR] Không đọc được {fpath}: {e}")
        return None


def extract_email_title(name_no_ext: str) -> str:
    """
    Lấy phần tựa đề từ tên file.
    Format: YYYYMMDD_Nguoi_gui_Ten_email  →  Ten_email (phần sau tên người gửi)
    Heuristic: bỏ phần YYYYMMDD_ đầu, giữ phần còn lại, thay _ bằng space.
    """
    # Bỏ prefix ngày YYYYMMDD_
    parts = name_no_ext.split("_", 1)
    if len(parts) == 2 and re.match(r'^\d{8}$', parts[0]):
        rest = parts[1]
    else:
        rest = name_no_ext
    return rest.replace("_", " ")


def count_words(text: str) -> int:
    return len([w for w in text.split() if w.strip()])


def summarize_with_ai(title: str, body: str) -> str | None:
    """Gửi (title + body) lên ChatGPT, trả về bản tóm tắt tiếng Việt."""
    word_count = count_words(body)
    if word_count < 50:
        print(f"  [SKIP] Bỏ qua '{title}' — chỉ có {word_count} từ.")
        return None

    prompt = (
        f"Tựa đề email: {title}\n\n"
        f"Nội dung:\n{body}"
    )

    try:
        response = client.chat.completions.create(
            model="gpt-4.1-mini-2025-04-14",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "Bạn là trợ lý chuyên tóm tắt email bằng tiếng Việt. "
                        "Tóm tắt trên 1500 chữ hoặc còn một nửa nội dung gốc (lấy giá trị lớn hơn). "
                        "Định dạng đầu ra: dòng đầu tiên là tựa đề bằng tiếng Việt (không có nhãn 'Tên bài:'), "
                        "tiếp theo là các đoạn văn xuôi, không dùng gạch đầu dòng hay liệt kê. "
                        "Không thêm lời giới thiệu hay giải thích."
                    )
                },
                {"role": "user", "content": prompt}
            ],
            temperature=1.0
        )
        summary = response.choices[0].message.content.strip()
        print(f"  [OK] '{title}' — gốc: {word_count} từ | tóm tắt: {count_words(summary)} từ")
        return summary
    except Exception as e:
        print(f"  [ERROR] AI lỗi với '{title}': {e}")
        return None


# ============================================================
# DOCX & EPUB EXPORT
# ============================================================

def build_docx(summaries: list[tuple[str, str]], docx_path: str) -> bool:
    """
    summaries: list of (title, summary_text)
    Tạo file DOCX với mỗi email là 1 heading + nội dung.
    """
    doc = Document()
    for title, content in summaries:
        lines = content.splitlines()
        # Dòng đầu tiên không rỗng → heading
        heading_set = False
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            if not heading_set:
                doc.add_heading(line, level=1)
                heading_set = True
            else:
                doc.add_paragraph(line)
        doc.add_paragraph()   # khoảng cách giữa các email

    try:
        os.makedirs(os.path.dirname(docx_path), exist_ok=True)
        doc.save(docx_path)
        print(f"\n[DOCX] Đã lưu: {docx_path}")
        return True
    except Exception as e:
        print(f"[ERROR] Không lưu được DOCX: {e}")
        return False


def build_epub(docx_path: str, epub_path: str):
    """Chuyển DOCX → EPUB với TOC."""
    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"[ERROR] Không đọc được DOCX để tạo EPUB: {e}")
        return

    book = epub.EpubBook()
    book.set_identifier("AI-Email-Summary")
    book.set_title("Tóm tắt Email AI")
    book.set_language("vi")

    chapters, toc = [], []
    chap_idx = 0
    cur_title, cur_content = None, []

    def flush_chapter():
        nonlocal chap_idx, cur_title, cur_content
        if cur_title is None:
            return
        chap_idx += 1
        chap = epub.EpubHtml(title=cur_title, file_name=f"chap_{chap_idx}.xhtml", lang="vi")
        html = f"<h1>{cur_title}</h1>"
        for p in cur_content:
            p_esc = p.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            html += f"<p>{p_esc}</p>"
        chap.content = html.encode("utf-8")
        book.add_item(chap)
        chapters.append(chap)
        toc.append(epub.Link(chap.file_name, cur_title, f"chap_{chap_idx}"))
        cur_title, cur_content = None, []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style.name == "Heading 1":
            flush_chapter()
            cur_title = text
        else:
            if cur_title:
                cur_content.append(text)

    flush_chapter()

    book.toc = toc
    book.spine = ["nav"] + chapters
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    try:
        epub.write_epub(epub_path, book, {})
        print(f"[EPUB] Đã lưu: {epub_path}")
    except Exception as e:
        print(f"[ERROR] Không tạo được EPUB: {e}")


# ============================================================
# HÀM CHÍNH: AI_sume
# ============================================================

def AI_sume(*args):
    """
    AI_sume(n, include_keywords, exclude_keywords)

    Tham số (tất cả đều tùy chọn, truyền theo vị trí):
      arg1 (int)          : số ngày gần nhất cần xử lý. Mặc định = 1 (hôm nay).
      arg2 (str/list/tuple): keyword(s) include — email có tựa đề chứa ít nhất 1 keyword này.
                             Bỏ qua / truyền None / () / "" → lấy tất cả.
      arg3 (str/list/tuple): keyword(s) exclude — email có tựa đề chứa keyword này sẽ bị loại.
                             Bỏ qua / truyền None / () / "" → không loại gì.

    Ví dụ:
      AI_sume()                              # hôm nay, tất cả email
      AI_sume(2)                             # 2 ngày gần nhất, tất cả email
      AI_sume(2, ("bloomberg", "burry"))     # 2 ngày, chỉ email có bloomberg hoặc burry
      AI_sume(3, None, ("Google",))          # 3 ngày, tất cả trừ email có 'google'
      AI_sume(1, "bloomberg", "google")      # hôm nay, có bloomberg, không có google
    """

    # ── Parse arguments ──────────────────────────────────────
    def _to_list(val) -> list[str]:
        """Chuẩn hóa keyword input → list[str]."""
        if val is None:
            return []
        if isinstance(val, str):
            return [val] if val.strip() else []
        if isinstance(val, (list, tuple)):
            return [str(k) for k in val if str(k).strip()]
        return []

    n            = int(args[0]) if len(args) > 0 and args[0] is not None else 1
    include_kws  = _to_list(args[1]) if len(args) > 1 else []
    exclude_kws  = _to_list(args[2]) if len(args) > 2 else []

    print("=" * 60)
    print(f"AI_sume  |  n={n}  |  include={include_kws}  |  exclude={exclude_kws}")
    print("=" * 60)

    # ── Lấy folder & file ────────────────────────────────────
    folders   = get_date_folders(n)
    all_files = get_email_files(folders)
    filtered  = filter_emails(all_files, include_kws, exclude_kws)

    print(f"Tổng email tìm thấy : {len(all_files)}")
    print(f"Sau khi lọc         : {len(filtered)}")

    if not filtered:
        print("Không có email nào phù hợp. Kết thúc.")
        return

    # ── Tóm tắt từng email ───────────────────────────────────
    summaries: list[tuple[str, str]] = []
    for fpath, name_no_ext in filtered:
        title = extract_email_title(name_no_ext)
        print(f"\n→ Đang xử lý: {title}")
        body = read_email_content(fpath)
        if body is None:
            continue
        summary = summarize_with_ai(title, body)
        if summary:
            summaries.append((title, summary))
        time.sleep(0.5)   # tránh rate-limit

    if not summaries:
        print("\nKhông có tóm tắt nào được tạo.")
        return

    # ── Xuất file ────────────────────────────────────────────
    today_str    = datetime.now().strftime("%Y-%m-%d")
    tag          = f"n{n}"
    if include_kws:
        tag += "_inc-" + "-".join(include_kws)[:30]
    if exclude_kws:
        tag += "_exc-" + "-".join(exclude_kws)[:30]

    out_folder   = os.path.join(OUTPUT_BASE_DIR, f"{today_str} - Summaries")
    os.makedirs(out_folder, exist_ok=True)

    docx_path    = os.path.join(out_folder, f"{today_str}_{tag}_summaries.docx")
    epub_path    = os.path.splitext(docx_path)[0] + ".epub"

    if build_docx(summaries, docx_path):
        build_epub(docx_path, epub_path)

    print(f"\n✅ Hoàn tất — {len(summaries)} email đã tóm tắt.")
    print(f"   DOCX : {docx_path}")
    print(f"   EPUB : {epub_path}")


# ============================================================
# CHẠY TRỰC TIẾP
# ============================================================
if __name__ == "__main__":
    # Ví dụ sử dụng — chỉnh lại theo nhu cầu:
    #AI_sume(2, ("bloomberg", "burry"), ())
    AI_sume(5, None, ("Google","burry","bloomberg"))
    # AI_sume()